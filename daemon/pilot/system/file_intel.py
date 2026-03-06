"""File Content Intelligence — understand file contents, not just metadata.

Parse PDFs, Word docs, Excel, images, and more. Extract text, summarize,
search within file contents, and analyze structure.
"""

from __future__ import annotations

import asyncio
import json
import logging
import mimetypes
import os
from pathlib import Path

logger = logging.getLogger("pilot.system.file_intel")


async def parse_file(file_path: str) -> str:
    """Auto-detect file type and extract text/data content."""
    p = Path(file_path)
    if not p.exists():
        return f"File not found: {file_path}"

    ext = p.suffix.lower()
    size = p.stat().st_size

    info = f"File: {p.name} ({size:,} bytes)\nType: {ext}\n\n"

    try:
        if ext == ".pdf":
            return info + await _parse_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return info + await _parse_docx(file_path)
        elif ext in (".xlsx", ".xls"):
            return info + await _parse_excel(file_path)
        elif ext == ".csv":
            return info + await _parse_csv(file_path)
        elif ext in (".json", ".jsonl"):
            return info + await _parse_json(file_path)
        elif ext in (".xml", ".html", ".htm"):
            return info + await _parse_xml(file_path)
        elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tiff"):
            return info + await _parse_image(file_path)
        elif ext in (".mp3", ".wav", ".flac", ".m4a", ".ogg"):
            return info + await _parse_audio_meta(file_path)
        elif ext in (".mp4", ".avi", ".mkv", ".mov", ".webm"):
            return info + await _parse_video_meta(file_path)
        elif ext in (".zip", ".tar", ".gz", ".7z", ".rar"):
            return info + await _parse_archive(file_path)
        elif ext in (".py", ".js", ".ts", ".java", ".cpp", ".c", ".h",
                      ".cs", ".go", ".rs", ".rb", ".php", ".swift", ".kt",
                      ".r", ".sql", ".sh", ".ps1", ".bat", ".toml", ".yaml",
                      ".yml", ".ini", ".cfg", ".conf", ".md", ".rst", ".txt",
                      ".log", ".env"):
            return info + await _parse_text(file_path)
        else:
            # Try as text
            try:
                return info + await _parse_text(file_path)
            except Exception:
                return info + f"Unable to parse file type: {ext}"
    except ImportError as e:
        return info + f"Missing dependency: {e}. Install it with pip."
    except Exception as e:
        return info + f"Parse error: {e}"


async def _parse_pdf(path: str) -> str:
    """Extract text from PDF."""
    # Try PyPDF2 / pypdf
    try:
        from pypdf import PdfReader

        def _do():
            reader = PdfReader(path)
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"--- Page {i+1} ---\n{text}")
            return f"Pages: {len(reader.pages)}\n\n" + "\n\n".join(pages)

        return await asyncio.to_thread(_do)
    except ImportError:
        pass

    # Try pdfplumber
    try:
        import pdfplumber

        def _do():
            with pdfplumber.open(path) as pdf:
                pages = []
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        pages.append(f"--- Page {i+1} ---\n{text}")
                return f"Pages: {len(pdf.pages)}\n\n" + "\n\n".join(pages)

        return await asyncio.to_thread(_do)
    except ImportError:
        raise ImportError("pypdf or pdfplumber (pip install pypdf pdfplumber)")


async def _parse_docx(path: str) -> str:
    """Extract text from Word documents."""
    from docx import Document

    def _do():
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            tables_text.append(f"--- Table {i+1} ---\n" + "\n".join(rows))

        output = "\n\n".join(paragraphs)
        if tables_text:
            output += "\n\n" + "\n\n".join(tables_text)
        return output

    return await asyncio.to_thread(_do)


async def _parse_excel(path: str) -> str:
    """Extract data from Excel files."""
    try:
        import openpyxl

        def _do():
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            output = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(max_row=100, values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    rows.append(" | ".join(cells))
                output.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
            return "\n\n".join(output)

        return await asyncio.to_thread(_do)
    except ImportError:
        raise ImportError("openpyxl (pip install openpyxl)")


async def _parse_csv(path: str) -> str:
    """Parse CSV file."""
    import csv

    def _do():
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            rows = []
            for i, row in enumerate(reader):
                if i >= 100:
                    rows.append(f"... ({i}+ rows total)")
                    break
                rows.append(" | ".join(row))
            return "\n".join(rows)

    return await asyncio.to_thread(_do)


async def _parse_json(path: str) -> str:
    """Parse JSON file."""
    text = Path(path).read_text("utf-8")
    data = json.loads(text)
    formatted = json.dumps(data, indent=2, ensure_ascii=False)
    if len(formatted) > 5000:
        formatted = formatted[:5000] + f"\n... ({len(text):,} chars total)"
    return formatted


async def _parse_xml(path: str) -> str:
    """Extract text from XML/HTML."""
    try:
        from bs4 import BeautifulSoup
        text = Path(path).read_text("utf-8", errors="replace")
        soup = BeautifulSoup(text, "html.parser")
        return soup.get_text(separator="\n", strip=True)[:5000]
    except ImportError:
        # Basic fallback
        import re
        text = Path(path).read_text("utf-8", errors="replace")
        clean = re.sub(r"<[^>]+>", "", text)
        return clean[:5000]


async def _parse_text(path: str) -> str:
    """Read plain text files."""
    text = Path(path).read_text("utf-8", errors="replace")
    lines = text.split("\n")
    line_count = len(lines)
    if line_count > 200:
        preview = "\n".join(lines[:200])
        return f"({line_count} lines total, showing first 200)\n\n{preview}"
    return text


async def _parse_image(path: str) -> str:
    """Get image metadata and optionally OCR text."""
    info_parts = []

    try:
        from PIL import Image
        from PIL.ExifTags import TAGS

        def _do():
            img = Image.open(path)
            parts = [
                f"Dimensions: {img.width}x{img.height}",
                f"Mode: {img.mode}",
                f"Format: {img.format}",
            ]
            # EXIF data
            exif_data = img.getexif()
            if exif_data:
                for tag_id, value in list(exif_data.items())[:20]:
                    tag = TAGS.get(tag_id, tag_id)
                    parts.append(f"  {tag}: {value}")
            return "\n".join(parts)

        info_parts.append(await asyncio.to_thread(_do))
    except ImportError:
        info_parts.append("(Install Pillow for image metadata: pip install Pillow)")

    return "\n".join(info_parts)


async def _parse_audio_meta(path: str) -> str:
    """Get audio file metadata."""
    try:
        from mutagen import File

        def _do():
            audio = File(path)
            if audio is None:
                return "Could not read audio metadata"
            parts = [f"Length: {audio.info.length:.1f}s"]
            if hasattr(audio.info, 'bitrate'):
                parts.append(f"Bitrate: {audio.info.bitrate // 1000}kbps")
            if hasattr(audio.info, 'sample_rate'):
                parts.append(f"Sample Rate: {audio.info.sample_rate}Hz")
            for key, value in list(audio.items())[:20]:
                parts.append(f"  {key}: {value}")
            return "\n".join(parts)

        return await asyncio.to_thread(_do)
    except ImportError:
        return "(Install mutagen for audio metadata: pip install mutagen)"


async def _parse_video_meta(path: str) -> str:
    """Get video file metadata."""
    import subprocess
    try:
        proc = await asyncio.create_subprocess_exec(
            "ffprobe", "-v", "quiet", "-print_format", "json",
            "-show_format", "-show_streams", path,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, _ = await proc.communicate()
        if stdout:
            data = json.loads(stdout.decode("utf-8"))
            fmt = data.get("format", {})
            parts = [
                f"Duration: {float(fmt.get('duration', 0)):.1f}s",
                f"Format: {fmt.get('format_long_name', 'unknown')}",
                f"Size: {int(fmt.get('size', 0)):,} bytes",
            ]
            for stream in data.get("streams", []):
                codec = stream.get("codec_type", "unknown")
                if codec == "video":
                    parts.append(f"Video: {stream.get('width')}x{stream.get('height')} {stream.get('codec_name')}")
                elif codec == "audio":
                    parts.append(f"Audio: {stream.get('codec_name')} {stream.get('sample_rate')}Hz")
            return "\n".join(parts)
    except FileNotFoundError:
        pass

    return f"Video file: {Path(path).stat().st_size:,} bytes (install ffprobe for metadata)"


async def _parse_archive(path: str) -> str:
    """List contents of an archive file."""
    ext = Path(path).suffix.lower()

    if ext == ".zip":
        import zipfile
        with zipfile.ZipFile(path, "r") as zf:
            items = zf.namelist()
            total = sum(i.file_size for i in zf.infolist())
            listing = "\n".join(items[:100])
            return f"ZIP: {len(items)} files, {total:,} bytes uncompressed\n\n{listing}"

    elif ext in (".tar", ".gz"):
        import tarfile
        with tarfile.open(path, "r:*") as tf:
            members = tf.getmembers()
            listing = "\n".join(m.name for m in members[:100])
            return f"TAR: {len(members)} entries\n\n{listing}"

    return f"Archive: {Path(path).stat().st_size:,} bytes"


async def search_file_contents(
    directory: str,
    search_text: str,
    pattern: str = "*.txt",
    max_results: int = 50,
) -> str:
    """Search for text within file contents (grep-like).

    Searches recursively through files matching the glob pattern.
    """
    results = []
    p = Path(directory)
    if not p.exists():
        return f"Directory not found: {directory}"

    def _do():
        count = 0
        for f in p.rglob(pattern):
            if count >= max_results:
                break
            if not f.is_file() or f.stat().st_size > 10_000_000:  # Skip >10MB
                continue
            try:
                for i, line in enumerate(f.open("r", encoding="utf-8", errors="replace"), 1):
                    if search_text.lower() in line.lower():
                        results.append({
                            "file": str(f),
                            "line": i,
                            "content": line.strip()[:200],
                        })
                        count += 1
                        if count >= max_results:
                            break
            except Exception:
                continue
        return results

    results = await asyncio.to_thread(_do)
    if not results:
        return f"No matches for '{search_text}' in {directory} ({pattern})"
    return json.dumps({"matches": results, "count": len(results)}, indent=2)
